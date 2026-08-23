#!/usr/bin/env python3
"""
Genera i tre allegati al Contratto DPA Art. 28 GDPR:
  All. A — Elenco attività di trattamento autorizzate
  All. B — Misure di sicurezza tecniche implementate
  All. C — Lista sub-responsabili autorizzati
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x08, 0x20, 0x50)
BLUE  = RGBColor(0x10, 0x78, 0xC0)
GRAY  = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREEN = RGBColor(0x16, 0x6B, 0x3A)
RED   = RGBColor(0x9B, 0x1C, 0x1C)
AMBER = RGBColor(0x92, 0x40, 0x0E)

# ── helpers condivisi ──────────────────────────────────────────────────────────

def set_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for s in doc.sections:
        s.top_margin    = Cm(top)
        s.bottom_margin = Cm(bottom)
        s.left_margin   = Cm(left)
        s.right_margin  = Cm(right)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='AAAAAA', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)

def add_hline(doc, color='1078C0'):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_para(doc, text='', bold=False, italic=False, size=10,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size      = Pt(size)
        run.font.color.rgb = color if color else BLACK
    return p

def add_heading(doc, text, level=1, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(12 if level == 1 else 10)
    run.font.color.rgb = NAVY if level == 1 else BLUE
    return p

def add_table(doc, headers, rows, col_widths=None, hdr_color='EEF6FB', hdr_border='1078C0'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = 'Table Grid'

    # header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, hdr_color)
        set_cell_border(cell, hdr_border, '6')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r   = p.add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY

    # rows
    for r_i, row_data in enumerate(rows):
        row = tbl.rows[r_i + 1]
        bg  = 'FFFFFF' if r_i % 2 == 0 else 'F4F9FE'
        for c_i, (txt, opt) in enumerate(row_data if isinstance(row_data[0], tuple) else [(x, {}) for x in row_data]):
            cell = row.cells[c_i]
            set_cell_bg(cell, opt.get('bg', bg))
            set_cell_border(cell, 'CCCCCC', '4')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r   = p.add_run(txt)
            r.bold           = opt.get('bold', False)
            r.font.size      = Pt(9)
            r.font.color.rgb = opt.get('color', GRAY)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return tbl

def doc_header(doc, titolo, sottotitolo, num_allegato):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('SINDROME RENU ITALIA APS')
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY

    for line in ['Associazione di Promozione Sociale — C.F./P.IVA 98020680157',
                 'Via Marina 6, 20121 Milano (MI) — www.sindromerenu.it']:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(line)
        r2.font.size = Pt(9); r2.font.color.rgb = GRAY

    add_hline(doc)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(6)
    p3.paragraph_format.space_after  = Pt(2)
    r3 = p3.add_run(f'ALLEGATO {num_allegato}')
    r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = BLUE

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(2)
    r4 = p4.add_run(titolo)
    r4.bold = True; r4.font.size = Pt(14); r4.font.color.rgb = NAVY

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_after = Pt(4)
    r5 = p5.add_run(sottotitolo)
    r5.font.size = Pt(10); r5.font.color.rgb = GRAY

    add_hline(doc)

    add_para(doc,
        'Allegato al Contratto DPA Art. 28 Reg. UE 2016/679 ("GDPR") tra '
        'Sindrome ReNU Italia APS (Titolare) e il Responsabile Tecnico del sito www.sindromerenu.it.',
        italic=True, size=9, color=GRAY, space_before=4, space_after=8)

def doc_footer_firme(doc):
    add_hline(doc)
    add_para(doc, 'Le Parti dichiarano di aver letto, compreso e approvato il presente allegato.',
             size=10, color=GRAY, space_before=6, space_after=10)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    dati = [
        ['Il Titolare del Trattamento', 'Il Responsabile del Trattamento'],
        ['Sindrome ReNU Italia APS\nStefania Rocca, Presidente', '[Nome e Cognome Responsabile Tecnico]'],
        ['Data: _____________________', 'Data: _____________________'],
        ['Firma: _____________________', 'Firma: _____________________'],
    ]
    for r_i, row in enumerate(tbl.rows):
        for c_i, cell in enumerate(row.cells):
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(dati[r_i][c_i])
            run.bold           = (r_i == 0)
            run.font.size      = Pt(10)
            run.font.color.rgb = NAVY if r_i == 0 else GRAY

# ═══════════════════════════════════════════════════════════════════════════════
# ALLEGATO A — Elenco attività di trattamento autorizzate
# ═══════════════════════════════════════════════════════════════════════════════
def crea_allegato_a():
    doc = Document()
    set_margins(doc)
    doc_header(doc,
               'ELENCO DETTAGLIATO DELLE ATTIVITÀ DI TRATTAMENTO AUTORIZZATE',
               'al Responsabile del Trattamento — Responsabile Tecnico del Sito',
               'A')

    add_para(doc,
        'Il presente allegato elenca tutte le attività di trattamento di dati personali che il Responsabile Tecnico '
        'è autorizzato a svolgere per conto di Sindrome ReNU Italia APS, con indicazione della base giuridica, '
        'delle categorie di dati, dei sistemi coinvolti e delle operazioni permesse. '
        'Qualsiasi attività non espressamente elencata richiede preventiva autorizzazione scritta del Titolare.',
        size=10, color=GRAY, space_after=8)

    # ── Sezione 1: Sviluppo e manutenzione ────────────────────────────────────
    add_heading(doc, '1. SVILUPPO E MANUTENZIONE DEL SITO WEB')
    add_para(doc, 'Piattaforma: Hono/TypeScript su Cloudflare Pages — https://sindromerenu-italia.pages.dev',
             size=9, italic=True, color=BLUE, space_after=6)

    hdrs = ['Attività', 'Sistemi coinvolti', 'Dati personali coinvolti', 'Operazioni permesse']
    rows_1 = [
        ['Sviluppo e aggiornamento codice sorgente applicazione',
         'Cloudflare Pages (src/index.tsx)',
         'Nessun dato personale — solo codice',
         'Scrittura, lettura, deploy'],
        ['Configurazione infrastruttura Cloudflare',
         'Cloudflare Dashboard, wrangler CLI',
         'Secret cifrati (ADMIN_SECRET, RESEND_API_KEY)',
         'Lettura, modifica, rotazione secret'],
        ['Gestione migrazioni database (schema)',
         'Cloudflare D1 — sindromerenu-db',
         'Nessun dato personale — solo struttura tabelle',
         'CREATE/ALTER/DROP TABLE — solo schema'],
        ['Deploy in produzione',
         'Cloudflare Pages deploy',
         'Nessun dato personale',
         'Upload bundle dist/'],
        ['Monitoraggio errori e log applicazione',
         'Cloudflare Dashboard, wrangler tail',
         'Possibili log con IP hashati o messaggi di errore',
         'Solo lettura — nessuna esportazione'],
    ]
    add_table(doc, hdrs, rows_1, col_widths=[4.5, 3.5, 4.5, 3.5])
    doc.add_paragraph()

    # ── Sezione 2: Gestione database dati personali ───────────────────────────
    add_heading(doc, '2. ACCESSO AL DATABASE — DATI PERSONALI DEGLI UTENTI')
    add_para(doc, 'Database: Cloudflare D1 — sindromerenu-db (ID: 8f747f7d-16ec-4ef8-ac1c-f4b5ce39a25d)',
             size=9, italic=True, color=BLUE, space_after=4)
    add_para(doc,
        '⚠  Alcune tabelle contengono dati di categoria speciale (Art. 9 GDPR): '
        'dati sanitari di minori (diagnosi genetica, varianti patogene, storia clinica). '
        'Per queste tabelle si applicano misure di sicurezza rafforzate.',
        size=9, color=AMBER, space_after=6)

    hdrs2 = ['Tabella DB', 'Dati personali contenuti', 'Categoria speciale', 'Operazioni permesse', 'Vincoli']
    rows_2 = [
        [('contatti',         {'bold': True, 'color': NAVY}),
         ('Nome, email, messaggio, hash IP (sha2:…), consenso v2.0, lingua, timestamp', {}),
         ('No', {'color': GREEN}),
         ('SELECT, INSERT — solo per manutenzione tecnica o richiesta Titolare', {}),
         ('Nessuna esportazione senza autorizzazione scritta', {})],

        [('lista_attesa',     {'bold': True, 'color': NAVY}),
         ('Nome, cognome, email, telefono, città, tipo, hash IP, consenso v2.0; dati minore: nome, anno nascita, patologia', {}),
         ('Parziale — dati minore', {'color': AMBER}),
         ('SELECT, INSERT, UPDATE, DELETE — su richiesta Titolare o per manutenzione', {}),
         ('Delete solo tramite pannello admin o richiesta esplicita Titolare', {})],

        [('adesioni',         {'bold': True, 'color': NAVY}),
         ('Nome, cognome, email, telefono, città, tipo socio, consenso, timestamp', {}),
         ('No', {'color': GREEN}),
         ('SELECT, INSERT, UPDATE, DELETE — su richiesta Titolare', {}),
         ('Nessuna comunicazione a terzi', {})],

        [('storie',           {'bold': True, 'color': NAVY}),
         ('Nome bambino, testo storia (5 lingue), URL foto, tipo, flag attiva', {}),
         ('SÌ — dati sanitari minori (Art. 9)', {'color': RED, 'bold': True}),
         ('SELECT, INSERT, UPDATE, DELETE — solo su istruzione esplicita del Titolare', {}),
         ('Trattamento subordinato a consenso esplicito acquisito dal Titolare; nessuna copia locale', {})],

        [('donazioni',        {'bold': True, 'color': NAVY}),
         ('Nome, email, importo, metodo, data, consenso', {}),
         ('No', {'color': GREEN}),
         ('SELECT — solo lettura', {}),
         ('Nessuna modifica o cancellazione autonoma; conservazione 10 anni per obbligo fiscale', {})],

        [('audit_log',        {'bold': True, 'color': NAVY}),
         ('Tipo azione, payload, hash IP, timestamp', {}),
         ('No', {'color': GREEN}),
         ('SELECT, INSERT (automatico da sistema)', {}),
         ('Solo lettura manuale; nessuna cancellazione prima dei 12 mesi', {})],

        [('faq / news / pubblicazioni / brochure / gallery / eventi / testi_ui', {'bold': True, 'color': NAVY}),
         ('Nessun dato personale — contenuti editoriali', {}),
         ('No', {'color': GREEN}),
         ('CRUD completo — gestione contenuti del sito', {}),
         ('Nessun vincolo specifico', {})],
    ]
    add_table(doc, hdrs2, rows_2, col_widths=[3.5, 4.5, 2.2, 4.0, 3.8])
    doc.add_paragraph()

    # ── Sezione 3: Pannello admin ─────────────────────────────────────────────
    add_heading(doc, '3. ACCESSO AL PANNELLO AMMINISTRATIVO  (/admin)')
    add_para(doc, 'Autenticazione: header X-Admin-Token verificato server-side contro secret ADMIN_SECRET cifrato su Cloudflare.',
             size=9, italic=True, color=BLUE, space_after=6)

    hdrs3 = ['Funzione admin', 'Endpoint API', 'Metodo', 'Dati coinvolti', 'Autorizzazione']
    rows_3 = [
        ['Statistiche generali',           '/api/admin/stats',              'GET',    'Conteggi aggregati — nessun dato personale',           'Sempre autorizzata'],
        ['Elenco contatti',                '/api/admin/contatti',           'GET',    'Tutti i dati tabella contatti',                        'Solo lettura'],
        ['Elenco lista attesa',            '/api/admin/lista-attesa',       'GET',    'Tutti i dati tabella lista_attesa',                    'Solo lettura'],
        ['Elenco adesioni',                '/api/admin/adesioni',           'GET',    'Tutti i dati tabella adesioni',                        'Solo lettura'],
        ['Elenco donazioni',               '/api/admin/donazioni',          'GET',    'Tutti i dati tabella donazioni',                       'Solo lettura'],
        ['Log operazioni admin',           '/api/admin/audit',              'GET',    'Audit log completo',                                   'Solo lettura'],
        ['Diritto all\'oblio',             '/api/admin/erasure/:email',     'DELETE', 'Pseudonimizzazione dati su TUTTE le tabelle per email', 'Solo su richiesta scritta Titolare'],
        ['Gestione storie',                '/api/admin/storie',             'CRUD',   'Dati sanitari minori (Art. 9)',                        'Solo su istruzione esplicita Titolare'],
        ['Gestione FAQ/news/eventi/brochure','/api/admin/faq|news|…',      'CRUD',   'Contenuti editoriali — nessun dato personale',         'Autonoma'],
        ['Gestione testi UI',              '/api/admin/testi_ui',           'GET/PUT','Traduzioni interfaccia — nessun dato personale',       'Autonoma'],
    ]
    add_table(doc, hdrs3, rows_3, col_widths=[3.8, 3.8, 1.5, 4.0, 4.0])
    doc.add_paragraph()

    # ── Sezione 4: Endpoint pubblici ──────────────────────────────────────────
    add_heading(doc, '4. ENDPOINT API PUBBLICI — RICEZIONE DATI DAGLI UTENTI')
    add_para(doc, 'Questi endpoint sono esposti su Internet e ricevono dati personali dagli utenti del sito.',
             size=10, color=GRAY, space_after=6)

    hdrs4 = ['Endpoint', 'Metodo', 'Dati ricevuti', 'Operazione DB', 'Misure di sicurezza applicate']
    rows_4 = [
        ['POST /api/lista-attesa', 'POST',
         'nome, cognome, email, telefono, città, tipo, consenso, dati minore',
         'INSERT in lista_attesa',
         'Sanitizzazione input (san()), validazione email, hash IP SHA-256, consenso GDPR v2.0 obbligatorio'],
        ['POST /api/contatti', 'POST',
         'nome, email, oggetto, messaggio, consenso',
         'INSERT in contatti',
         'Sanitizzazione input, validazione email, hash IP SHA-256, consenso GDPR v2.0 obbligatorio'],
        ['GET /api/storie', 'GET',
         'Nessun dato ricevuto — solo lettura pubblica',
         'SELECT su storie (attive=1)',
         'Nessun dato personale in output oltre a quelli pubblicati con consenso'],
        ['GET /api/faq|news|pubblicazioni|gallery|eventi', 'GET',
         'Nessun dato personale',
         'SELECT su tabelle contenuti',
         'Nessuna misura specifica richiesta'],
    ]
    add_table(doc, hdrs4, rows_4, col_widths=[3.5, 1.8, 4.0, 3.2, 5.5])
    doc.add_paragraph()

    # ── Sezione 5: Attività NON autorizzate ───────────────────────────────────
    add_heading(doc, '5. ATTIVITÀ ESPRESSAMENTE NON AUTORIZZATE')
    add_para(doc,
        'Il Responsabile del Trattamento NON è autorizzato a svolgere le seguenti attività '
        'senza esplicita e preventiva autorizzazione scritta del Titolare:',
        size=10, color=GRAY, space_after=6)

    non_auth = [
        'Esportare, copiare o trasferire a terzi dati personali degli utenti in qualsiasi forma.',
        'Comunicare credenziali di accesso al pannello admin o ai servizi Cloudflare a soggetti non autorizzati.',
        'Attivare l\'endpoint sendEmail() / Resend API senza preventiva comunicazione al Titolare e aggiornamento informativa.',
        'Modificare o cancellare dati della tabella donazioni (conservazione obbligatoria 10 anni per obbligo fiscale).',
        'Implementare strumenti di analytics, tracking o profilazione degli utenti non previsti dall\'informativa vigente.',
        'Aggiungere sub-responsabili (altri fornitori tecnici) che trattino dati personali senza autorizzazione scritta del Titolare.',
        'Trattare dati sanitari dei minori (tabella storie) per finalità diverse dalla pubblicazione sul sito con consenso acquisito.',
        'Conservare copie locali del database su dispositivi personali o server non autorizzati dal Titolare.',
    ]
    for item in non_auth:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    doc.add_paragraph()
    doc_footer_firme(doc)

    path = '/home/user/webapp-renu/Allegato_A_Attivita_Trattamento_Autorizzate.docx'
    doc.save(path)
    print(f'✅  {path}')
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# ALLEGATO B — Misure di sicurezza tecniche implementate
# ═══════════════════════════════════════════════════════════════════════════════
def crea_allegato_b():
    doc = Document()
    set_margins(doc)
    doc_header(doc,
               'MISURE DI SICUREZZA TECNICHE IMPLEMENTATE',
               'Documento Tecnico — Art. 32 Reg. UE 2016/679 (GDPR)',
               'B')

    add_para(doc,
        'Il presente allegato descrive le misure tecniche e organizzative adottate ai sensi dell\'Art. 32 GDPR '
        'per garantire un livello di sicurezza adeguato al rischio dei trattamenti effettuati tramite '
        'il sito www.sindromerenu.it. Le misure sono verificabili direttamente nel codice sorgente '
        '(src/index.tsx) e nella configurazione dell\'infrastruttura Cloudflare.',
        size=10, color=GRAY, space_after=8)

    # ── 1. Trasmissione dati ──────────────────────────────────────────────────
    add_heading(doc, '1. TRASMISSIONE DATI — HTTPS / TLS')

    hdrs = ['Misura', 'Implementazione', 'Verifica']
    rows = [
        ['Cifratura in transito HTTPS/TLS',
         'Tutta la rete Cloudflare gestisce il TLS automaticamente. '
         'Il sito è raggiungibile esclusivamente via HTTPS; '
         'le richieste HTTP sono reindirizzate a HTTPS da Cloudflare.',
         'https://sindromerenu-italia.pages.dev — certificato TLS valido (Cloudflare Universal SSL)'],
        ['HSTS (HTTP Strict Transport Security)',
         'Abilitato automaticamente da Cloudflare Pages per tutti i siti .pages.dev',
         'Header Strict-Transport-Security presente nelle risposte HTTP'],
        ['Cifratura dati in transito database',
         'Cloudflare D1 utilizza connessioni TLS interne per la comunicazione tra Worker e database SQLite distribuito.',
         'Documentazione Cloudflare D1 Security'],
    ]
    add_table(doc, hdrs, rows, col_widths=[4.0, 7.0, 5.5])
    doc.add_paragraph()

    # ── 2. Pseudonimizzazione IP ──────────────────────────────────────────────
    add_heading(doc, '2. PSEUDONIMIZZAZIONE INDIRIZZI IP (SHA-256)')

    add_para(doc, 'Riferimento codice sorgente: src/index.tsx, funzione hashIPAsync() — riga ~5148',
             size=9, italic=True, color=BLUE, space_after=4)

    box = doc.add_paragraph()
    box.paragraph_format.left_indent  = Cm(1.0)
    box.paragraph_format.space_before = Pt(2)
    box.paragraph_format.space_after  = Pt(6)
    r_box = box.add_run(
        'async function hashIPAsync(ip: string): Promise<string> {\n'
        '  const encoder = new TextEncoder()\n'
        '  const data = encoder.encode(\'renu-ip-salt-2026:\' + ip)\n'
        '  const hashBuffer = await crypto.subtle.digest(\'SHA-256\', data)\n'
        '  const hashArray = Array.from(new Uint8Array(hashBuffer))\n'
        '  return \'sha2:\' + hashArray.map(b => b.toString(16).padStart(2,\'0\')).join(\'\').slice(0,16)\n'
        '}'
    )
    r_box.font.name = 'Courier New'; r_box.font.size = Pt(8.5); r_box.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    rows_ip = [
        ['Algoritmo',         'SHA-256 (Web Crypto API — crypto.subtle.digest)',
                              'Standard crittografico — non reversibile con mezzi computazionali attuali'],
        ['Salt applicato',    '"renu-ip-salt-2026:" prefissato all\'IP prima del digest',
                              'Impedisce attacchi rainbow table sull\'intero spazio di indirizzi IPv4'],
        ['Output',            'Prefisso "sha2:" + primi 16 caratteri hex del digest (es. sha2:a3f2b1c8d9e0f4a1)',
                              'Identificatore pseudonimo non riconducibile all\'IP originale'],
        ['Conservazione IP',  'L\'indirizzo IP originale NON viene mai scritto nel database.',
                              'La funzione getClientIP() legge l\'header CF-Connecting-IP ma '
                              'il valore grezzo non raggiunge mai lo storage'],
        ['Fallback',          'In caso di errore crypto.subtle: valore "ip:unavailable" — nessun fallback non sicuro',
                              'Nessun uso di Math.imul() o altri metodi non crittografici'],
        ['Tabelle interessate','contatti (campo ip_hash), lista_attesa (campo ip_hash)',
                              'Tutte le altre tabelle non registrano dati di rete'],
    ]
    add_table(doc, ['Caratteristica', 'Dettaglio implementazione', 'Note di sicurezza'],
              rows_ip, col_widths=[3.5, 7.0, 6.0])
    doc.add_paragraph()

    # ── 3. Accesso admin ──────────────────────────────────────────────────────
    add_heading(doc, '3. CONTROLLO ACCESSO PANNELLO AMMINISTRATIVO')
    add_para(doc, 'Riferimento: src/index.tsx, funzione requireAdmin() — riga ~6054',
             size=9, italic=True, color=BLUE, space_after=4)

    rows_admin = [
        ['Meccanismo autenticazione',
         'Token Bearer nel header HTTP X-Admin-Token, verificato server-side ad ogni richiesta API.',
         'Nessuna sessione persistente; ogni chiamata API è autenticata indipendentemente.'],
        ['Storage del secret',
         'ADMIN_SECRET cifrato come Cloudflare Workers Secret (equivalente a secret manager). '
         'Non è mai presente nel codice sorgente (src/index.tsx), nei file di configurazione '
         'o nel repository git.',
         'npx wrangler secret list — verifica presenza senza esporre il valore'],
        ['Fallback di sviluppo',
         'In assenza del secret Cloudflare (ambiente locale): valore di default "renu-admin-2026". '
         'Questo valore non è mai in produzione perché il secret è sempre configurato.',
         'Il fallback è accettabile solo in ambiente di sviluppo locale'],
        ['Risposta su accesso non autorizzato',
         'HTTP 401 Unauthorized con corpo JSON {"error": "Non autorizzato"}. '
         'Nessuna informazione aggiuntiva esposta.',
         'Nessuna rivelazione della struttura interna degli endpoint'],
        ['Endpoint protetti',
         'Tutti gli endpoint /api/admin/* sono protetti da requireAdmin(). '
         'Totale: 25+ endpoint verificati.',
         'Nessun endpoint admin è accessibile senza autenticazione'],
        ['Pannello HTML admin',
         'La pagina /admin è HTML statico; il token viene inserito dall\'operatore '
         'e trasmesso via JavaScript al momento delle chiamate API.',
         'Il token non è mai nel markup HTML — solo in memoria JavaScript lato client'],
    ]
    add_table(doc, ['Misura', 'Implementazione', 'Note'], rows_admin, col_widths=[3.5, 7.5, 5.5])
    doc.add_paragraph()

    # ── 4. Consenso GDPR ─────────────────────────────────────────────────────
    add_heading(doc, '4. GESTIONE CONSENSO GDPR — VERSIONING')

    rows_cons = [
        ['Versione consenso',
         'Entrambi i form (lista attesa e contatti) registrano la versione del testo '
         'di consenso al momento dell\'invio (campo gdpr_consent_version = "2.0").',
         'Consente di dimostrare il consenso specifico acquisito per ogni versione del testo'],
        ['Testo consenso',
         'Il testo completo non è archiviato nel database; la versione "2.0" corrisponde '
         'al testo presente nel codice sorgente al momento del deploy.',
         'In caso di aggiornamento testo: incrementare versione e documentare in CHANGELOG'],
        ['Timestamp consenso',
         'Campo created_at DATETIME DEFAULT CURRENT_TIMESTAMP in tutte le tabelle.',
         'Tracciabilità temporale del consenso acquisito'],
        ['Obbligatorietà',
         'I form non possono essere inviati senza spuntare il checkbox GDPR (validazione client + server).',
         'Nessun inserimento in DB senza consenso registrato'],
        ['Banner cookie',
         'Funzione acceptCookies() con storage in localStorage (fallback: document.cookie). '
         'Chiave: renu_cookie_consent, valore: "1". '
         'Banner mostrato 800ms dopo DOMContentLoaded se consenso non ancora registrato.',
         'Nessun cookie di profilazione impostato prima del consenso'],
    ]
    add_table(doc, ['Misura', 'Implementazione', 'Note'], rows_cons, col_widths=[3.5, 7.5, 5.5])
    doc.add_paragraph()

    # ── 5. Diritto all'oblio ──────────────────────────────────────────────────
    add_heading(doc, '5. DIRITTO ALL\'OBLIO — PSEUDONIMIZZAZIONE AUTOMATICA (ART. 17 GDPR)')
    add_para(doc, 'Endpoint: DELETE /api/admin/erasure/:email — riga ~6141 src/index.tsx',
             size=9, italic=True, color=BLUE, space_after=6)

    rows_oblio = [
        ['Trigger',
         'Richiesta DELETE autenticata a /api/admin/erasure/{email_interessato}',
         'Solo il Responsabile Tecnico può eseguirla tramite pannello admin'],
        ['Tabelle interessate',
         'contatti, lista_attesa, adesioni, donazioni, storie — tutte le tabelle con dati personali',
         'Operazione multi-tabella atomica'],
        ['Metodo di cancellazione',
         'PSEUDONIMIZZAZIONE: i campi personali sono sovrascritti con valori non identificabili '
         '(es. nome → "[cancellato]", email → "[gdpr-erasure-{timestamp}]"). '
         'Il record viene mantenuto per integrità referenziale ma reso non identificabile.',
         'Non viene eseguita una DELETE fisica per preservare l\'audit trail e l\'integrità del DB'],
        ['Tracciabilità',
         'Ogni esecuzione dell\'erasure viene registrata in audit_log con: '
         'tipo="erasure", email_target (hashata), timestamp, numero record modificati.',
         'Dimostrabilità dell\'esercizio del diritto all\'oblio'],
        ['Eccezioni',
         'I dati di donazione soggetti a obbligo fiscale (10 anni) sono pseudonimizzati '
         'ma il record contabile è mantenuto.',
         'Conformità Art. 17 par. 3.b GDPR — eccezione per obbligo legale'],
    ]
    add_table(doc, ['Aspetto', 'Implementazione', 'Note'], rows_oblio, col_widths=[3.0, 8.0, 5.5])
    doc.add_paragraph()

    # ── 6. Sanitizzazione input ───────────────────────────────────────────────
    add_heading(doc, '6. VALIDAZIONE E SANITIZZAZIONE INPUT')
    add_para(doc, 'Funzione san() — src/index.tsx riga ~5142',
             size=9, italic=True, color=BLUE, space_after=4)

    rows_san = [
        ['Sanitizzazione input',
         'Funzione san(v, n) applicata a tutti i campi prima dell\'inserimento in DB: '
         'casting a stringa, trim(), troncamento a n caratteri (default 500).',
         'Prevenzione SQL injection e overflow campi'],
        ['Validazione email',
         'Funzione validEmail() con regex RFC-compliant applicata prima di ogni INSERT.',
         'Rifiuto di indirizzi malformati prima che raggiungano il database'],
        ['Prepared statements',
         'Tutte le query SQL usano il metodo .bind() di Cloudflare D1 (parametri bind). '
         'Nessuna concatenazione di stringhe nelle query.',
         'Protezione completa da SQL injection'],
        ['Validazione server-side',
         'La validazione è eseguita nel Worker server-side (src/index.tsx), '
         'non solo lato client. Il client-side è solo UX.',
         'Impossibile bypassare la validazione manipolando il browser'],
    ]
    add_table(doc, ['Misura', 'Implementazione', 'Note'], rows_san, col_widths=[3.5, 7.5, 5.5])
    doc.add_paragraph()

    # ── 7. Infrastruttura e certificazioni ────────────────────────────────────
    add_heading(doc, '7. INFRASTRUTTURA — CERTIFICAZIONI E STANDARD')

    rows_infra = [
        ['Cloudflare Pages', 'Edge computing globale', 'ISO 27001, SOC 2 Type II, PCI DSS Level 1', 'https://www.cloudflare.com/trust-hub/'],
        ['Cloudflare D1',    'Database SQLite distribuito', 'ISO 27001 (eredita da Cloudflare infra)', 'Cloudflare Trust Hub'],
        ['Cloudflare Workers','Runtime serverless edge', 'ISO 27001, SOC 2', 'Cloudflare Trust Hub'],
        ['TLS/SSL',          'Trasmissione cifrata', 'Cloudflare Universal SSL — cert. automatici', 'TLS 1.2 / 1.3'],
    ]
    add_table(doc,
              ['Componente', 'Ruolo', 'Certificazioni', 'Riferimento'],
              rows_infra, col_widths=[3.5, 3.5, 5.5, 4.0])
    doc.add_paragraph()

    # ── 8. Audit log ──────────────────────────────────────────────────────────
    add_heading(doc, '8. AUDIT LOG — TRACCIABILITÀ OPERAZIONI')

    rows_audit = [
        ['Tabella audit_log', 'Registra tutte le operazioni admin: tipo azione, payload (dati coinvolti), '
         'IP hashato dell\'operatore, timestamp.',
         'Schema: id, action_type, action_data (JSON), ip_hash, created_at'],
        ['Operazioni tracciate', 'Inserimenti dati, modifiche contenuti, cancellazioni, '
         'esercizio diritto all\'oblio, accessi al pannello.',
         'Ogni chiamata agli endpoint /api/admin/* genera un record'],
        ['Conservazione',
         '12 mesi dalla registrazione (poi cancellazione automatica non implementata — '
         'da gestire manualmente o con cron job futuro).',
         'Da implementare: pulizia automatica log >12 mesi'],
        ['Accesso audit log',
         'Solo lettura tramite pannello admin (/api/admin/audit — GET). '
         'Nessuna modifica o cancellazione permessa tramite API.',
         'Integrità del log garantita dalla sola possibilità di lettura via API'],
    ]
    add_table(doc, ['Aspetto', 'Implementazione', 'Note'], rows_audit, col_widths=[3.5, 8.0, 5.0])
    doc.add_paragraph()

    # ── 9. Misure organizzative ───────────────────────────────────────────────
    add_heading(doc, '9. MISURE ORGANIZZATIVE')
    misure_org = [
        'Le credenziali di accesso al pannello admin e ai servizi Cloudflare sono note solo al Responsabile Tecnico e non vengono condivise senza autorizzazione del Titolare.',
        'Il codice sorgente è versionato su repository git privato; nessun secret o dato personale è presente nella cronologia dei commit.',
        'Le operazioni di accesso ai dati personali vengono eseguite solo su specifica richiesta del Titolare o per necessità tecniche documentate.',
        'In caso di data breach: notifica al Titolare entro 24 ore, con indicazione di tipo violazione, dati coinvolti, numero stimato di interessati e misure adottate.',
        'Il Responsabile Tecnico non conserva copie locali dei dati del database su dispositivi personali.',
    ]
    for m in misure_org:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(m)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    doc.add_paragraph()
    doc_footer_firme(doc)

    path = '/home/user/webapp-renu/Allegato_B_Misure_Sicurezza_Tecniche.docx'
    doc.save(path)
    print(f'✅  {path}')
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# ALLEGATO C — Lista sub-responsabili autorizzati
# ═══════════════════════════════════════════════════════════════════════════════
def crea_allegato_c():
    doc = Document()
    set_margins(doc)
    doc_header(doc,
               'LISTA SUB-RESPONSABILI AUTORIZZATI',
               'Fornitori tecnici che trattano dati per conto del Responsabile Tecnico',
               'C')

    add_para(doc,
        'Ai sensi dell\'Art. 28 par. 2 GDPR, il Responsabile del Trattamento può ricorrere a un altro '
        'responsabile del trattamento (sub-responsabile) solo previa autorizzazione scritta del Titolare. '
        'Il presente allegato elenca i sub-responsabili già autorizzati dal Titolare in quanto componenti '
        'tecniche essenziali dell\'infrastruttura del sito www.sindromerenu.it, nonché i soggetti che '
        'agiscono come titolari autonomi del trattamento nell\'erogazione dei propri servizi tecnici.',
        size=10, color=GRAY, space_after=6)

    # ── Nota Cloudflare doppio ruolo ──────────────────────────────────────────
    add_para(doc,
        'NOTA — DOPPIO RUOLO DI CLOUDFLARE: Cloudflare, Inc. compare in due sezioni distinte del presente '
        'allegato perché esercita due ruoli giuridici separati e indipendenti: '
        '(1) responsabile del trattamento per i servizi di hosting/database che gestisce su istruzione '
        'dell\'Associazione (Pages, Workers, D1); '
        '(2) titolare autonomo del trattamento per i dati tecnici dei visitatori (IP, log di rete) '
        'che elabora per proprie finalità di sicurezza della rete, indipendentemente dall\'Associazione. '
        'Fonte: Cloudflare Privacy Policy, Sez. 6.',
        size=9, italic=True, color=AMBER, space_after=10)

    # ══════════════════════════════════════════════════════════════════════════
    # SEZIONE 1 — SUB-RESPONSABILI ART. 28
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'SEZIONE 1 — SUB-RESPONSABILI DEL TRATTAMENTO (Art. 28 GDPR)')
    add_para(doc,
        'I seguenti soggetti trattano dati personali degli utenti del sito esclusivamente per conto '
        'e su istruzione del Titolare o del Responsabile Tecnico, nell\'ambito dei servizi descritti.',
        size=10, color=GRAY, space_after=8)

    # ── 1a. Cloudflare come responsabile ─────────────────────────────────────
    add_heading(doc, '1a. CLOUDFLARE, INC. — Hosting, Workers e Database D1', level=2)

    def add_campo(doc, label, val):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(label + '  ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        r2 = p.add_run(val)
        r2.font.size = Pt(10); r2.font.color.rgb = GRAY

    campi_cf = [
        ('Ragione sociale:',          'Cloudflare, Inc. (società madre: Plus Five Five, Inc.)'),
        ('Sede legale:',              '101 Townsend Street, San Francisco, CA 94107, USA'),
        ('Ruolo in questa sezione:',  'Sub-responsabile del trattamento (Art. 28 GDPR) — '
                                      'per i soli servizi di hosting, runtime e database elencati sotto'),
        ('Servizi autorizzati:',
         '1) Cloudflare Pages — hosting e distribuzione dell\'applicazione web\n'
         '2) Cloudflare Workers — runtime serverless per il codice server-side\n'
         '3) Cloudflare D1 — database SQLite distribuito (sindromerenu-db, '
         'ID: 8f747f7d-16ec-4ef8-ac1c-f4b5ce39a25d)\n'
         '4) Cloudflare Workers Secrets — storage cifrato dei secret ADMIN_SECRET e RESEND_API_KEY'),
        ('Dati personali trattati\ncome sub-responsabile:',
         'Dati degli utenti archiviati nel database D1:\n'
         '  • tabella contatti: nome, email, messaggio, hash IP, consenso\n'
         '  • tabella lista_attesa: nome, cognome, email, telefono, città, dati minore\n'
         '  • tabella adesioni: nome, cognome, email, telefono\n'
         '  • tabella donazioni: nome, email, importo\n'
         '  • tabella storie: nome bambino, storia, dati sanitari (cat. speciale Art. 9)\n'
         'I secret cifrati (ADMIN_SECRET, RESEND_API_KEY) non sono dati personali degli interessati.'),
        ('Modalità di accesso\nai dati:',
         'Cloudflare accede ai dati D1 solo per erogare il servizio di database distribuito '
         '(storage, replication, backup). Non analizza né utilizza il contenuto per finalità proprie.'),
        ('Base giuridica\ntrasferimento USA:',
         'EU-US Data Privacy Framework (DPF) — certificazione Cloudflare attiva\n'
         'Decisione di adeguatezza Commissione Europea n. 2023/1795 del 10 luglio 2023\n'
         'Base: Art. 45 GDPR (paese con livello adeguato) — non richiede SCCs\n'
         'Verifica certificazione: https://www.dataprivacyframework.gov/participant/4585'),
        ('DPA:',
         'Cloudflare Data Processing Addendum (DPA) — incorporato nei termini di servizio\n'
         'Accettato al momento della creazione dell\'account Cloudflare\n'
         'Testo: https://www.cloudflare.com/cloudflare-customer-dpa/'),
        ('Certificazioni sicurezza:',
         'ISO/IEC 27001:2013 · SOC 2 Type II · PCI DSS Level 1 · CSA STAR Level 1\n'
         'Riferimento: https://www.cloudflare.com/trust-hub/'),
        ('Privacy Policy:',           'https://www.cloudflare.com/privacypolicy/'),
        ('Progetto Cloudflare Pages:', 'sindromerenu-italia — https://sindromerenu-italia.pages.dev'),
        ('Autorizzato dal:',
         'Titolare — Sindrome ReNU Italia APS — alla firma del presente Accordo'),
    ]
    for label, val in campi_cf:
        add_campo(doc, label, val)

    add_hline(doc, color='CCCCCC')

    # ── 1b. Resend ────────────────────────────────────────────────────────────
    add_heading(doc, '1b. RESEND INC. (Plus Five Five, Inc.) — Email Transazionali  ⚠ NON ANCORA ATTIVO', level=2)

    add_para(doc,
        '⚠  Il servizio Resend è predisposto nel codice (funzione sendEmail()) ma NON è attualmente '
        'collegato ad alcun endpoint attivo. Nessun dato personale viene trasmesso a Resend. '
        'L\'autorizzazione sotto è condizionata: diventerà operativa solo al completamento '
        'degli adempimenti elencati nel campo "Condizioni di attivazione".',
        size=10, italic=True, color=AMBER, space_before=4, space_after=6)

    campi_resend = [
        ('Ragione sociale:',          'Plus Five Five, Inc. (nome commerciale: Resend)'),
        ('Sede legale:',              '2261 Market Street #5039, San Francisco, CA 94114, USA'),
        ('Ruolo GDPR:',
         'Sub-responsabile del trattamento (Art. 28 GDPR) per i dati dei destinatari delle email '
         '(Customer Data nella terminologia Resend).\n'
         'Titolare autonomo per i dati dell\'account Resend dell\'Associazione '
         '(Company Account Data) e per i log d\'uso del servizio (Company Usage Data).\n'
         'Fonte: Resend DPA, Sez. 2.1 e Sez. 9.'),
        ('Servizi previsti\n(quando attivato):',
         'Invio di email transazionali agli utenti del sito: '
         'notifiche di ricezione form, conferme iscrizione, risposte automatiche'),
        ('Dati personali\ntrattati come\nsub-responsabile:',
         'Dati dei destinatari delle email (utenti del sito che compilano i form):\n'
         '  • indirizzo email del destinatario\n'
         '  • nome del destinatario (se incluso nel messaggio)\n'
         '  • contenuto del messaggio email\n'
         '  • metadati di consegna (timestamp, IP di consegna)\n'
         'Fonte: Resend DPA Exhibit A — "Categories of personal data transferred"'),
        ('Base giuridica\ntrasferimento USA:',
         'EU-US Data Privacy Framework (DPF) — certificazione Resend attiva dal 13 marzo 2025\n'
         'Decisione di adeguatezza Commissione Europea n. 2023/1795 del 10 luglio 2023\n'
         'Base: Art. 45 GDPR — non richiede SCCs come meccanismo principale\n'
         'Il DPA Resend prevede anche SCCs (Dec. CE 2021/914/UE) come meccanismo di fallback '
         'qualora il DPF cessasse di essere valido (Resend DPA, Sez. 6.2)\n'
         'Verifica certificazione DPF: https://www.dataprivacyframework.gov/\n'
         'Annuncio certificazione: https://resend.com/changelog/data-privacy-framework-certification'),
        ('DPA:',
         'Resend Data Processing Addendum — incorporato e accettato automaticamente '
         'all\'accettazione dei Terms of Service Resend.\n'
         'Non è necessaria una firma separata: il DPA diventa vincolante '
         'al momento della creazione dell\'account Resend.\n'
         'Testo completo: https://resend.com/legal/dpa'),
        ('Pagina GDPR Resend:',       'https://resend.com/security/gdpr'),
        ('Privacy Policy:',           'https://resend.com/legal/privacy-policy'),
        ('Sub-processori Resend:',    'https://resend.com/legal/subprocessors\n'
                                      '(Resend notifica modifiche ai sub-processori con 14 giorni di preavviso)'),
        ('Condizioni di attivazione:',
         'Prima di attivare la funzione sendEmail() nel codice, il Responsabile Tecnico deve:\n'
         '  1) Creare un account Resend (accettando automaticamente il DPA)\n'
         '  2) Comunicare al Titolare e al DPO (dpo@sindromerenu.it) l\'attivazione\n'
         '  3) Aggiornare l\'informativa privacy del sito (sezione 8) con Resend come sub-responsabile\n'
         '  4) Aggiornare il presente allegato con la data di attivazione effettiva\n'
         '  5) Verificare che la RESEND_API_KEY sia configurata come secret Cloudflare'),
        ('Status:',
         '❌ NON ATTIVO — nessun dato personale trasmesso a Resend alla data del presente allegato'),
        ('Autorizzato dal:',
         'Titolare — autorizzazione condizionata agli adempimenti sopra elencati'),
    ]
    for label, val in campi_resend:
        add_campo(doc, label, val)

    add_hline(doc, color='CCCCCC')

    # ══════════════════════════════════════════════════════════════════════════
    # SEZIONE 2 — TITOLARI AUTONOMI
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'SEZIONE 2 — TITOLARI AUTONOMI DEL TRATTAMENTO (non sub-responsabili)')
    add_para(doc,
        'I seguenti soggetti NON sono sub-responsabili del trattamento ai sensi dell\'Art. 28 GDPR. '
        'Agiscono come titolari autonomi (independent controllers) per i dati tecnici che elaborano '
        'nell\'ambito dei propri servizi infrastrutturali, per proprie finalità e secondo proprie policy. '
        'Il Titolare e il Responsabile Tecnico non impartiscono istruzioni su tali trattamenti '
        'e non possono limitarli: è sufficiente menzionarli nell\'informativa privacy con rinvio '
        'alle rispettive privacy policy.',
        size=10, color=GRAY, space_after=6)

    hdrs_aut = ['Soggetto', 'Servizio', 'Dati elaborati come titolare autonomo',
                'Base trasferimento extra-UE', 'Privacy Policy']
    rows_aut = [
        ['Cloudflare, Inc.\n(ruolo autonomo —\ndistinto dalla Sez. 1a)',
         'CDN globale, protezione DDoS, sicurezza rete — elaborazione delle richieste HTTP '
         'dei visitatori prima che raggiungano il server',
         'Indirizzo IP visitatori, header HTTP, User-Agent, dati di routing e timing —'
         ' elaborati per proprie finalità di sicurezza e ottimizzazione rete.\n'
         'Fonte: Cloudflare Privacy Policy, Sez. 6:\n'
         '"Cloudflare is a data controller for the personal information collected '
         'from all categories of data subjects listed above"',
         'EU-US DPF\n(Dec. CE 2023/1795)\nArt. 45 GDPR',
         'cloudflare.com/\nprivacypolicy/'],
        ['jsDelivr CDN\n(ProspectOne\nSp. z o.o.)\nPolonia — UE',
         'Distribuzione file CSS/JS statici:\nTailwind CSS e FontAwesome '
         '(caricati dal browser del visitatore al primo accesso al sito)',
         'Indirizzo IP del visitatore trasmesso al server CDN jsDelivr '
         'al momento del caricamento degli asset CSS/JS.\n'
         'Nessun cookie di profilazione impostato da jsDelivr.',
         'Trattamento interamente in UE — nessun trasferimento extra-UE\n'
         'Non richiede SCCs né DPF',
         'jsdelivr.com/terms/\nprivacy-policy-\njsdelivr-net'],
    ]
    add_table(doc, hdrs_aut, rows_aut, col_widths=[3.2, 3.8, 5.5, 2.8, 3.2])
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SEZIONE 3 — PROCEDURA
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'SEZIONE 3 — PROCEDURA PER L\'AGGIUNTA DI NUOVI SUB-RESPONSABILI')
    add_para(doc,
        'Prima di ricorrere a qualsiasi nuovo sub-responsabile non elencato nel presente allegato, '
        'il Responsabile del Trattamento deve:',
        size=10, color=GRAY, space_after=4)

    steps = [
        'Comunicare per iscritto al Titolare (info@sindromerenu.it) e al DPO (dpo@sindromerenu.it): '
        'nome e sede del soggetto, ruolo previsto, categorie di dati che tratterà, '
        'garanzie GDPR offerte (DPA, certificazioni, base trasferimento extra-UE).',
        'Attendere autorizzazione scritta del Titolare prima di procedere.',
        'Verificare che il sub-responsabile offra garanzie adeguate ai sensi dell\'Art. 28 par. 1 GDPR '
        '(DPA disponibile e firmato/accettato, certificazioni di sicurezza, '
        'base trasferimento extra-UE valida: DPF preferibile, SCCs come alternativa).',
        'Aggiornare il presente allegato con le informazioni complete del nuovo sub-responsabile '
        'e la data di autorizzazione.',
        'Comunicare l\'aggiunta agli interessati tramite aggiornamento dell\'informativa privacy '
        'sul sito (sezione 8 — tabella responsabili del trattamento).',
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f'{i}.  ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = BLUE
        r2 = p.add_run(step)
        r2.font.size = Pt(10); r2.font.color.rgb = GRAY

    # ══════════════════════════════════════════════════════════════════════════
    # SEZIONE 4 — RIEPILOGO
    # ══════════════════════════════════════════════════════════════════════════
    add_para(doc, '', space_after=4)
    add_heading(doc, 'SEZIONE 4 — RIEPILOGO STATO AUTORIZZAZIONI', space_before=6)

    hdrs_rie = ['Soggetto', 'Ruolo GDPR', 'Base trasferimento USA', 'Status', 'Note']
    rows_rie = [
        [('Cloudflare, Inc.',       {'bold': True, 'color': NAVY}),
         ('Sub-responsabile\n(Pages, Workers, D1, Secrets)', {}),
         ('EU-US DPF\n(Art. 45 GDPR)', {'color': GREEN}),
         ('✅ Autorizzato',         {'color': GREEN, 'bold': True}),
         ('DPA accettato nei ToS Cloudflare', {})],

        [('Cloudflare, Inc.',       {'bold': True, 'color': NAVY}),
         ('Titolare autonomo\n(dati rete visitatori)', {}),
         ('EU-US DPF\n(Art. 45 GDPR)', {'color': GREEN}),
         ('ℹ️ Titolare autonomo\nnessuna nomina richiesta', {'color': BLUE}),
         ('Menzionare nell\'informativa\ncon link privacy policy', {})],

        [('Resend Inc.',            {'bold': True, 'color': NAVY}),
         ('Sub-responsabile\n(email transazionali)', {}),
         ('EU-US DPF primario\nSCCs come fallback', {'color': GREEN}),
         ('⚠️ Condizionato\nnon ancora attivo', {'color': AMBER, 'bold': True}),
         ('DPA accettato nei ToS\nResend all\'attivazione', {})],

        [('jsDelivr / ProspectOne', {'bold': True, 'color': NAVY}),
         ('Titolare autonomo\n(CDN CSS/JS — UE)', {}),
         ('Nessuno — UE', {'color': GREEN}),
         ('ℹ️ Titolare autonomo\nnessuna nomina richiesta', {'color': BLUE}),
         ('Menzionare nell\'informativa\ncon link privacy policy', {})],
    ]
    add_table(doc, hdrs_rie, rows_rie, col_widths=[3.2, 3.5, 3.3, 3.3, 4.2])
    doc.add_paragraph()

    doc_footer_firme(doc)

    path = '/home/user/webapp-renu/Allegato_C_Sub_Responsabili_Autorizzati.docx'
    doc.save(path)
    print(f'✅  {path}')
    return path


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    pa = crea_allegato_a()
    pb = crea_allegato_b()
    pc = crea_allegato_c()
    print(f'\n📎 Tre allegati generati:\n  {pa}\n  {pb}\n  {pc}')
