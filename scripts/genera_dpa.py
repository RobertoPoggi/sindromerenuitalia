#!/usr/bin/env python3
"""
Genera i due documenti DOCX:
  1. DPA Art. 28 GDPR - Contratto Responsabile Tecnico
  2. Lettera all'Avv. Conti (testo mail)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colori brand ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x08, 0x20, 0x50)   # #082050
BLUE   = RGBColor(0x10, 0x78, 0xC0)   # #1078C0
GRAY   = RGBColor(0x55, 0x55, 0x55)
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xEE, 0xF6, 0xFB)   # sfondo intestazioni tabella

# ── Helper: imposta shading cella ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: bordo cella ────────────────────────────────────────────────────────
def set_cell_border(cell, color='AAAAAA', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ── Helper: paragrafo con testo ────────────────────────────────────────────────
def add_para(doc, text='', bold=False, italic=False, size=10,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color if color else BLACK
    return p

# ── Helper: intestazione sezione ──────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(12 if level == 1 else 10)
    run.font.color.rgb = NAVY if level == 1 else BLUE
    return p

# ── Helper: riga con label+valore ─────────────────────────────────────────────
def add_field(doc, label, value='_______________'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(value)
    r2.bold = False
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY

# ── Helper: tabella styled ─────────────────────────────────────────────────────
def add_styled_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    # Intestazione
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, 'EEF6FB')
        set_cell_border(cell, '1078C0', '6')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY

    # Righe dati
    for r_i, row_data in enumerate(rows):
        row = tbl.rows[r_i + 1]
        bg = 'FFFFFF' if r_i % 2 == 0 else 'F7FBFF'
        for c_i, cell_text in enumerate(row_data):
            cell = row.cells[c_i]
            set_cell_bg(cell, bg)
            set_cell_border(cell, 'CCCCCC', '4')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY

    # Larghezze colonne
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl

# ── Helper: riga firma ─────────────────────────────────────────────────────────
def add_firma_table(doc, label1, label2):
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = [(0, label1), (0, label2)]
    for row_i, row in enumerate(tbl.rows):
        for c_i, cell in enumerate(row.cells):
            p = cell.paragraphs[0]
            if row_i == 0:
                titles = [label1, label2]
                run = p.add_run(titles[c_i])
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = NAVY
            elif row_i == 1:
                subs = [
                    'Sindrome ReNU Italia APS\nStefania Rocca, Presidente',
                    '[Nome e Cognome Responsabile Tecnico]'
                ]
                run = p.add_run(subs[c_i])
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY
            elif row_i == 2:
                run = p.add_run('Data: _____________________')
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY
            elif row_i == 3:
                run = p.add_run('Firma: _____________________')
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY
    return tbl

# ── Linea separatrice ─────────────────────────────────────────────────────────
def add_hline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1078C0')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1 — CONTRATTO DPA ART. 28 GDPR
# ═══════════════════════════════════════════════════════════════════════════════
def crea_dpa():
    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Intestazione documento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('SINDROME RENU ITALIA APS')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run('Associazione di Promozione Sociale — C.F./P.IVA 98020680157')
    r2.font.size = Pt(9); r2.font.color.rgb = GRAY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run('Via Marina 6, 20121 Milano (MI) | info@sindromerenu.it | www.sindromerenu.it')
    r3.font.size = Pt(9); r3.font.color.rgb = GRAY

    add_hline(doc)

    # Titolo
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(10)
    pt.paragraph_format.space_after  = Pt(4)
    rt = pt.add_run('ACCORDO SUL TRATTAMENTO DEI DATI PERSONALI')
    rt.bold = True; rt.font.size = Pt(14); rt.font.color.rgb = NAVY

    pt2 = doc.add_paragraph()
    pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt2.paragraph_format.space_after = Pt(2)
    rt2 = pt2.add_run('Data Processing Agreement — Art. 28 Reg. UE 2016/679 "GDPR"')
    rt2.bold = True; rt2.font.size = Pt(11); rt2.font.color.rgb = BLUE

    add_hline(doc)

    # ── PARTI ─────────────────────────────────────────────────────────────────
    add_heading(doc, 'PARTI DEL CONTRATTO')

    add_para(doc, 'TRA', bold=True, color=NAVY, space_before=6, space_after=2)
    add_field(doc, 'Denominazione:', 'Sindrome ReNU Italia APS')
    add_field(doc, 'Forma giuridica:', 'Associazione di Promozione Sociale iscritta al RUNTS dal 28/04/2025')
    add_field(doc, 'C.F./P.IVA:', '98020680157')
    add_field(doc, 'Sede legale:', 'Via Marina 6, 20121 Milano (MI)')
    add_field(doc, 'PEC:', 'sindromerenuitalia@legalmail.it')
    add_field(doc, 'Rappresentata da:', 'Stefania Rocca, Presidente pro tempore')
    p_label = doc.add_paragraph()
    r_label = p_label.add_run('di seguito denominata "Titolare del Trattamento" o "Titolare"')
    r_label.italic = True; r_label.font.size = Pt(10); r_label.font.color.rgb = GRAY
    p_label.paragraph_format.space_after = Pt(8)

    add_para(doc, 'E', bold=True, color=NAVY, space_before=2, space_after=2)
    add_field(doc, 'Nome e Cognome:',  '[___________________________]')
    add_field(doc, 'Codice Fiscale:',  '[___________________________]')
    add_field(doc, 'Residente/sede:',  '[___________________________]')
    add_field(doc, 'Email:',           '[___________________________]')
    add_field(doc, 'PEC (se disp.):',  '[___________________________]')
    p_label2 = doc.add_paragraph()
    r_label2 = p_label2.add_run('di seguito denominato/a "Responsabile del Trattamento" o "Responsabile Tecnico"')
    r_label2.italic = True; r_label2.font.size = Pt(10); r_label2.font.color.rgb = GRAY
    p_label2.paragraph_format.space_after = Pt(8)

    add_hline(doc)

    # ── PREMESSE ──────────────────────────────────────────────────────────────
    add_heading(doc, 'PREMESSE')
    premesse = [
        ('A.', 'Il Titolare è una Associazione di Promozione Sociale che gestisce il sito web istituzionale www.sindromerenu.it, tramite il quale raccoglie e tratta dati personali degli utenti (moduli di contatto, lista attesa, adesioni, donazioni, storie delle famiglie).'),
        ('B.', 'Il Responsabile del Trattamento ha sviluppato e gestisce tecnicamente il sito, con accesso diretto al database e al pannello amministrativo, ed elabora pertanto dati personali per conto del Titolare nell\'ambito del mandato tecnico conferitogli.'),
        ('C.', 'L\'Art. 28 del Regolamento UE 2016/679 (GDPR) richiede che il trattamento da parte di un Responsabile sia disciplinato da un contratto scritto vincolante che stabilisca oggetto, durata, natura, finalità del trattamento, le categorie di dati e gli obblighi e diritti del Titolare.'),
        ('D.', 'Le Parti concordano di formalizzare con il presente Accordo i termini e le condizioni applicabili al trattamento dei dati personali svolto dal Responsabile del Trattamento per conto del Titolare.'),
    ]
    for letter, text in premesse:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(letter + '  ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = BLUE
        r2 = p.add_run(text)
        r2.font.size = Pt(10); r2.font.color.rgb = GRAY

    add_hline(doc)

    # ── ART. 1 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 1 — OGGETTO E DURATA')
    add_para(doc, '1.1  Il presente Accordo disciplina il trattamento dei dati personali che il Responsabile del Trattamento effettua per conto del Titolare nell\'ambito delle seguenti attività:', size=10, color=GRAY, space_after=3)
    attivita = [
        'sviluppo, manutenzione e aggiornamento del sito web www.sindromerenu.it;',
        'gestione del database Cloudflare D1 (sindromerenu-db) contenente i dati degli utenti;',
        'accesso al pannello amministrativo /admin per la gestione operativa dei dati;',
        'implementazione di misure di sicurezza tecniche (cifratura, hashing SHA-256, pseudonimizzazione);',
        'esecuzione del diritto all\'oblio su richiesta del Titolare (endpoint DELETE /api/admin/erasure/:email);',
        'gestione dei secret Cloudflare Workers (ADMIN_SECRET, RESEND_API_KEY) e delle credenziali di accesso ai servizi terzi.',
    ]
    for a in attivita:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(a)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    add_para(doc, '1.2  Il presente Accordo ha efficacia dalla data di sottoscrizione e rimane valido per tutta la durata del rapporto di collaborazione tecnica, salvo disdetta scritta con preavviso di 30 giorni.', size=10, color=GRAY, space_before=4)

    # ── ART. 2 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 2 — ISTRUZIONI DEL TITOLARE')
    istr = [
        '2.1  Il Responsabile del Trattamento tratta i dati personali esclusivamente su istruzione documentata del Titolare, salvo diversi obblighi di legge.',
        '2.2  Le istruzioni del Titolare comprendono, a titolo esemplificativo: eseguire operazioni di lettura, modifica, cancellazione o pseudonimizzazione dei dati su specifica richiesta; implementare nuove funzionalità che comportano trattamento di dati personali solo previa autorizzazione scritta del Titolare; non cedere, comunicare o trasferire a terzi i dati personali senza esplicita autorizzazione.',
        '2.3  Qualora il Responsabile del Trattamento ritenga che un\'istruzione violi il GDPR o altra normativa applicabile, ne informa immediatamente il Titolare per iscritto, indicando i profili di criticità.',
    ]
    for t in istr:
        add_para(doc, t, size=10, color=GRAY, space_after=4)

    # ── ART. 3 — TABELLA DATI ─────────────────────────────────────────────────
    add_heading(doc, 'ART. 3 — CATEGORIE DI DATI E INTERESSATI')
    add_para(doc, '3.1  Le categorie di dati personali trattate dal Responsabile del Trattamento e le relative categorie di interessati sono le seguenti:', size=10, color=GRAY, space_after=6)

    hdrs3 = ['Categoria', 'Dati trattati', 'Interessati']
    rows3 = [
        ['Richieste informazioni',       'Nome, email, messaggio, hash IP (SHA-256)',                          'Utenti del sito'],
        ['Lista attesa / Adesioni',      'Nome, cognome, email, telefono, città, dati minore, consenso',       'Famiglie con bambini ReNU'],
        ['Storie famiglie ⚠ Art. 9',    'Nome bambino, storia, fotografia, diagnosi genetica (dati sanitari)','Famiglie — minori con Sindrome ReNU'],
        ['Donazioni',                    'Nome, email, importo, metodo, data, consenso',                       'Donatori'],
        ['Log amministrativi',           'Azioni, timestamp, hash IP',                                        'Utenti e operatori admin'],
    ]
    add_styled_table(doc, hdrs3, rows3, col_widths=[4.5, 7.5, 5.0])

    add_para(doc, '⚠  Dati di categoria speciale (Art. 9 GDPR): il database contiene dati sanitari relativi a minori (diagnosi genetica, varianti patogene, storia clinica). Il Responsabile del Trattamento è tenuto ad applicare misure di sicurezza rafforzate per questa categoria.', size=9, color=RGBColor(0x92, 0x40, 0x0E), space_before=6, space_after=4)

    # ── ART. 4 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 4 — OBBLIGHI DEL RESPONSABILE DEL TRATTAMENTO')

    obblighi = [
        ('4.1  Riservatezza',
         'Garantire che le persone autorizzate al trattamento abbiano assunto impegni di riservatezza o siano soggette a obbligo legale di riservatezza. Non divulgare, cedere o comunicare dati personali a terzi non autorizzati.'),
        ('4.2  Sicurezza tecnica',
         'Implementare e mantenere misure tecniche e organizzative adeguate ai sensi dell\'Art. 32 GDPR, tra cui: trasmissione esclusivamente tramite HTTPS/TLS; hashing crittografico degli indirizzi IP (SHA-256 con salt, non reversibile); accesso al pannello admin protetto da token segreto cifrato (mai nel codice sorgente); conservazione sicura delle credenziali di accesso ai servizi Cloudflare; nessuna copia locale dei dati del database su dispositivi non protetti.'),
        ('4.3  Sub-responsabili',
         'Non ricorrere a ulteriori responsabili del trattamento (sub-responsabili) senza previa autorizzazione scritta specifica del Titolare. L\'infrastruttura Cloudflare Pages / D1 è già autorizzata dal Titolare come componente tecnica del servizio.'),
        ('4.4  Assistenza al Titolare',
         'Assistere il Titolare nell\'adempimento degli obblighi di cui agli Artt. 32-36 GDPR e nell\'esercizio dei diritti degli interessati (accesso, rettifica, cancellazione, portabilità) fornendo tempestivamente le informazioni e gli strumenti tecnici necessari.'),
        ('4.5  Notifica violazioni',
         'Notificare al Titolare, senza ingiustificato ritardo (e comunque entro 24 ore dalla scoperta), qualsiasi violazione dei dati personali (data breach) di cui venga a conoscenza, fornendo le informazioni di cui all\'Art. 33 par. 3 GDPR.'),
        ('4.6  Cancellazione o restituzione',
         'Al termine del rapporto di collaborazione, su richiesta del Titolare, cancellare o restituire tutti i dati personali trattati e cancellare le copie esistenti, salvo diverso obbligo di legge.'),
        ('4.7  Registro dei trattamenti',
         'Mantenere il registro delle attività di trattamento svolte per conto del Titolare, ai sensi dell\'Art. 30 par. 2 GDPR.'),
        ('4.8  Cooperazione e audit',
         'Mettere a disposizione del Titolare e del DPO tutte le informazioni necessarie a dimostrare il rispetto degli obblighi del presente Accordo e consentire, facilitandole, le attività di revisione e ispezione da parte del Titolare o di un soggetto da esso incaricato.'),
    ]
    for title, body in obblighi:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(title + '  ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        r2 = p.add_run(body)
        r2.font.size = Pt(10); r2.font.color.rgb = GRAY

    # ── ART. 5 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 5 — DIRITTI E OBBLIGHI DEL TITOLARE')
    art5 = [
        '5.1  Il Titolare fornisce al Responsabile del Trattamento istruzioni chiare, documentate e conformi al GDPR.',
        '5.2  Il Titolare è responsabile della liceità dei trattamenti, dell\'acquisizione dei consensi e dell\'adeguatezza dell\'informativa agli interessati.',
        '5.3  Il Titolare notifica al Responsabile del Trattamento qualsiasi richiesta di esercizio dei diritti da parte degli interessati che richieda intervento tecnico.',
    ]
    for t in art5:
        add_para(doc, t, size=10, color=GRAY, space_after=4)

    # ── ART. 6 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 6 — TRASFERIMENTI INTERNAZIONALI')
    art6 = [
        '6.1  Il Responsabile del Trattamento non trasferisce i dati personali verso paesi terzi al di fuori dello Spazio Economico Europeo senza previa autorizzazione scritta del Titolare.',
        '6.2  L\'infrastruttura Cloudflare comporta l\'elaborazione di dati tecnici (IP, log) su server distribuiti globalmente, inclusi USA. Questo trattamento è effettuato da Cloudflare come titolare autonomo del trattamento (independent controller) nell\'ambito del proprio servizio di rete, ed è legittimato dall\'adesione di Cloudflare all\'EU-US Data Privacy Framework (Dec. adeguatezza CE 2023/1795). Il Responsabile Tecnico non è parte di questo rapporto.',
    ]
    for t in art6:
        add_para(doc, t, size=10, color=GRAY, space_after=4)

    # ── ART. 7 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 7 — COMPENSO E NATURA DEL RAPPORTO')
    add_para(doc, '7.1  Le condizioni economiche della collaborazione tecnica sono disciplinate da separato accordo tra le Parti.', size=10, color=GRAY, space_after=4)
    add_para(doc, '7.2  Il presente Accordo non costituisce contratto di lavoro subordinato né alcun altro rapporto che esuli dall\'incarico tecnico descritto.', size=10, color=GRAY, space_after=4)

    # ── ART. 8 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 8 — RESPONSABILITÀ')
    add_para(doc, '8.1  Il Responsabile del Trattamento risponde nei confronti del Titolare per i danni derivanti dal mancato rispetto degli obblighi del presente Accordo.', size=10, color=GRAY, space_after=4)
    add_para(doc, '8.2  Ciascuna Parte risponde nei confronti degli interessati e dell\'Autorità di controllo per i propri inadempimenti ai sensi degli Artt. 82-83 GDPR.', size=10, color=GRAY, space_after=4)

    # ── ART. 9 ────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 9 — LEGGE APPLICABILE E FORO COMPETENTE')
    add_para(doc, 'Il presente Accordo è regolato dal diritto italiano e dal Regolamento UE 2016/679. Per ogni controversia è competente il Foro di Milano.', size=10, color=GRAY, space_after=6)

    add_hline(doc)

    # ── FIRME ─────────────────────────────────────────────────────────────────
    add_heading(doc, 'ART. 10 — FIRME')
    add_para(doc, 'Letto, compreso e approvato in ogni sua parte. Firmato in duplice originale.', size=10, color=GRAY, space_after=10)

    add_firma_table(doc, 'Il Titolare del Trattamento', 'Il Responsabile del Trattamento')

    doc.add_paragraph()
    add_hline(doc)

    # ── ALLEGATI ──────────────────────────────────────────────────────────────
    add_heading(doc, 'ALLEGATI')
    allegati = [
        'All. A — Elenco dettagliato delle attività di trattamento autorizzate',
        'All. B — Misure di sicurezza tecniche implementate (documento tecnico)',
        'All. C — Lista sub-responsabili autorizzati (Cloudflare Pages / D1)',
    ]
    for a in allegati:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(a)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    path = '/home/user/webapp-renu/DPA_Responsabile_Tecnico_Art28_GDPR.docx'
    doc.save(path)
    print(f'✅ Salvato: {path}')
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2 — LETTERA ALL'AVV. CONTI
# ═══════════════════════════════════════════════════════════════════════════════
def crea_lettera():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Intestazione mittente
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Sindrome ReNU Italia APS')
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY

    for line in [
        'Via Marina 6 — 20121 Milano (MI)',
        'C.F. / P.IVA: 98020680157',
        'info@sindromerenu.it  |  sindromerenuitalia@legalmail.it',
        'www.sindromerenu.it',
    ]:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(line)
        r2.font.size = Pt(9); r2.font.color.rgb = GRAY

    add_hline(doc)

    # Data e destinatario
    add_para(doc, 'Milano, [data]', size=10, color=GRAY, space_before=6, space_after=8)

    p_dest = doc.add_paragraph()
    p_dest.paragraph_format.space_after = Pt(2)
    r_dest = p_dest.add_run('Egr. Avv. Francesco Conti')
    r_dest.bold = True; r_dest.font.size = Pt(11); r_dest.font.color.rgb = NAVY

    for line in [
        'Responsabile della Protezione dei Dati (DPO)',
        'nominato ex Art. 37 GDPR — Sindrome ReNU Italia APS',
        'dpo@sindromerenu.it',
    ]:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(1)
        r3 = p3.add_run(line)
        r3.font.size = Pt(10); r3.font.color.rgb = GRAY

    doc.add_paragraph()
    p_ogg = doc.add_paragraph()
    p_ogg.paragraph_format.space_after = Pt(8)
    r_o1 = p_ogg.add_run('Oggetto: ')
    r_o1.bold = True; r_o1.font.size = Pt(10); r_o1.font.color.rgb = NAVY
    r_o2 = p_ogg.add_run('Ricognizione dei trattamenti dati del sito www.sindromerenu.it — documentazione tecnica per adempimenti GDPR')
    r_o2.font.size = Pt(10); r_o2.font.color.rgb = BLACK

    # Saluto apertura
    add_para(doc, 'Gentile Avv. Conti,', bold=True, size=10, color=NAVY, space_after=4)
    add_para(doc,
        'in seguito alla Sua nomina a Responsabile della Protezione dei Dati (DPO) di questa Associazione, '
        'Le trasmettiamo la ricognizione completa dei trattamenti di dati personali effettuati tramite il sito web '
        'istituzionale www.sindromerenu.it, comprensiva della descrizione tecnica dell\'infrastruttura, dei soggetti '
        'coinvolti e delle misure di sicurezza adottate.',
        size=10, color=GRAY, space_after=8)

    add_hline(doc)

    # ── 1. TITOLARE ───────────────────────────────────────────────────────────
    add_heading(doc, '1. TITOLARE DEL TRATTAMENTO')
    campi_titolare = [
        ('Denominazione:',    'Sindrome ReNU Italia APS'),
        ('Natura giuridica:', 'Associazione di Promozione Sociale iscritta al RUNTS dal 28/04/2025'),
        ('C.F. / P.IVA:',    '98020680157'),
        ('Sede legale:',      'Via Marina 6, 20121 Milano (MI)'),
        ('PEC:',              'sindromerenuitalia@legalmail.it'),
        ('Tel. Segreteria:',  '+39 327 763 4894'),
        ('Tel. Presidenza:',  '+39 335 730 1206  (Stefania Rocca)'),
    ]
    for label, val in campi_titolare:
        add_field(doc, label, val)

    doc.add_paragraph()

    # ── 2. RESPONSABILE TECNICO ───────────────────────────────────────────────
    add_heading(doc, '2. RESPONSABILE TECNICO DEL SITO (Art. 28 GDPR)')
    add_para(doc,
        'Il sito è stato sviluppato e viene gestito tecnicamente da un Responsabile Tecnico esterno, nominato '
        'Responsabile del Trattamento ai sensi dell\'Art. 28 GDPR, con il quale l\'Associazione ha stipulato '
        '(o dovrà stipulare) il relativo contratto scritto.',
        size=10, color=GRAY, space_after=4)
    compiti = [
        'ha sviluppato interamente il sito su piattaforma Hono/TypeScript deployato su Cloudflare Pages;',
        'gestisce il database Cloudflare D1 (sindromerenu-db) contenente i dati personali degli utenti;',
        'ha accesso al pannello amministrativo /admin tramite token segreto cifrato (ADMIN_SECRET);',
        'può eseguire operazioni di lettura, modifica ed esercizio del diritto all\'oblio (endpoint DELETE /api/admin/erasure/:email);',
        'gestisce i secret Cloudflare Workers (ADMIN_SECRET, RESEND_API_KEY) tramite il pannello Cloudflare.',
    ]
    for c in compiti:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(c)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    add_para(doc,
        'Si allega la bozza di contratto DPA Art. 28 GDPR da formalizzare tra l\'Associazione e il Responsabile Tecnico. '
        'Si chiede all\'Avv. Conti di verificarne la conformità prima della sottoscrizione.',
        size=10, italic=True, color=BLUE, space_before=6, space_after=8)

    add_hline(doc)

    # ── 3. CATEGORIE DI DATI ──────────────────────────────────────────────────
    add_heading(doc, '3. CATEGORIE DI DATI TRATTATI E FINALITÀ')

    sezioni_dati = [
        ('3a. Form di contatto  (/api/contatti) — Tabella DB: contatti',
         [('Dati:', 'Nome, email, messaggio, consenso GDPR v2.0, hash IP SHA-256, lingua'),
          ('Base giuridica:', 'Consenso (Art. 6.1.a GDPR)'),
          ('Conservazione:', '2 anni'),
          ('Nota:', 'L\'IP non viene mai conservato in chiaro; è sostituito da un hash SHA-256 non reversibile con salt fisso.')]),
        ('3b. Lista attesa / Pre-iscrizione  (/api/lista-attesa) — Tabella DB: lista_attesa',
         [('Dati:', 'Nome, cognome, email, telefono, città, tipo, consenso GDPR v2.0, hash IP SHA-256; dati minore: nome, anno nascita, patologia'),
          ('Base giuridica:', 'Consenso / esecuzione di un contratto (Art. 6.1.a-b GDPR)'),
          ('Conservazione:', 'Durata rapporto associativo + 5 anni')]),
        ('3c. Storie delle famiglie  (/api/storie) — Tabella DB: storie',
         [('Dati:', 'Nome bambino, storia (testo), fotografia, dati sanitari (diagnosi, variante genetica) — CATEGORIA SPECIALE Art. 9 GDPR'),
          ('Base giuridica:', 'Consenso esplicito (Art. 9.2.a GDPR)'),
          ('Conservazione:', 'Fino a revoca del consenso')]),
        ('3d. Donazioni  (/api/donazioni) — Tabella DB: donazioni',
         [('Dati:', 'Nome, email, importo, metodo, data, consenso'),
          ('Base giuridica:', 'Obbligo legale — conservazione fiscale (Art. 6.1.c GDPR)'),
          ('Conservazione:', '10 anni (obbligo fiscale ex D.P.R. 600/1973)'),
          ('Nota:', 'L\'endpoint POST /api/donazioni è predisposto nel codice ma non ancora attivo; le donazioni tramite bonifico non transitano attualmente dal sito.')]),
        ('3e. Adesioni associate  (/api/adesioni) — Tabella DB: adesioni',
         [('Dati:', 'Nome, cognome, email, telefono, città, tipo socio, consenso'),
          ('Base giuridica:', 'Esecuzione di contratto (Art. 6.1.b GDPR)')]),
        ('3f. Log di navigazione — Tabella DB: audit_log',
         [('Dati:', 'Tipo azione, dati coinvolti, IP hashato (SHA-256), timestamp'),
          ('Finalità:', 'Sicurezza informatica, tracciabilità operazioni admin'),
          ('Conservazione:', '12 mesi')]),
    ]

    for titolo_sez, campi_sez in sezioni_dati:
        p_sez = doc.add_paragraph()
        p_sez.paragraph_format.space_before = Pt(8)
        p_sez.paragraph_format.space_after  = Pt(3)
        r_sez = p_sez.add_run(titolo_sez)
        r_sez.bold = True; r_sez.font.size = Pt(10); r_sez.font.color.rgb = BLUE
        for lbl, val in campi_sez:
            p_f = doc.add_paragraph()
            p_f.paragraph_format.space_before = Pt(1)
            p_f.paragraph_format.space_after  = Pt(1)
            p_f.paragraph_format.left_indent  = Cm(0.8)
            r_l = p_f.add_run(lbl + '  ')
            r_l.bold = True; r_l.font.size = Pt(10); r_l.font.color.rgb = NAVY
            r_v = p_f.add_run(val)
            r_v.font.size = Pt(10); r_v.font.color.rgb = GRAY

    add_hline(doc)

    # ── 4. DIRITTO ALL'OBLIO ──────────────────────────────────────────────────
    add_heading(doc, '4. DIRITTO ALL\'OBLIO — Implementazione tecnica (Art. 17 GDPR)')
    add_para(doc,
        'È implementato nel pannello admin l\'endpoint tecnico:  DELETE /api/admin/erasure/:email',
        bold=True, size=10, color=NAVY, space_after=3)
    add_para(doc,
        'L\'operazione esegue la pseudonimizzazione di tutti i dati personali riconducibili all\'indirizzo email specificato, '
        'su tutte le tabelle del database (contatti, lista_attesa, adesioni, donazioni, storie). '
        'La pseudonimizzazione sostituisce i dati con valori non identificabili e registra l\'operazione nell\'audit_log '
        'per tracciabilità. L\'accesso a questo endpoint è riservato al solo Responsabile Tecnico tramite token admin cifrato.',
        size=10, color=GRAY, space_after=8)

    add_hline(doc)

    # ── 5. PANNELLO ADMIN ─────────────────────────────────────────────────────
    add_heading(doc, '5. PANNELLO AMMINISTRATIVO')
    add_para(doc,
        'Il pannello /admin è accessibile esclusivamente tramite token segreto (X-Admin-Token) e consente:',
        size=10, color=GRAY, space_after=6)

    hdrs_admin = ['Funzione', 'Tabella', 'Tipo accesso']
    rows_admin = [
        ['Visualizzazione contatti',         'contatti',    'Solo lettura'],
        ['Gestione lista attesa',             'lista_attesa','Lettura + eliminazione'],
        ['Gestione adesioni',                 'adesioni',    'Lettura + eliminazione'],
        ['Gestione donazioni',                'donazioni',   'Solo lettura'],
        ['Visualizzazione audit log',         'audit_log',   'Solo lettura'],
        ['Gestione storie famiglie',          'storie',      'Lettura + modifica + eliminazione'],
        ['Gestione contenuti (FAQ, news…)',   'varie',       'CRUD completo'],
        ['Esercizio diritto all\'oblio',      'tutte',       'Pseudonimizzazione per email'],
    ]
    add_styled_table(doc, hdrs_admin, rows_admin, col_widths=[7.0, 4.0, 5.0])

    doc.add_paragraph()
    add_hline(doc)

    # ── 6. INFRASTRUTTURA ─────────────────────────────────────────────────────
    add_heading(doc, '6. INFRASTRUTTURA TECNICA E SOGGETTI TERZI')

    add_para(doc, '6a. Cloudflare, Inc. — CDN e infrastruttura di rete',
             bold=True, size=10, color=BLUE, space_before=6, space_after=2)
    add_para(doc, '101 Townsend St, San Francisco, CA 94107, USA',
             italic=True, size=9, color=GRAY, space_after=4)
    add_para(doc,
        'Qualificazione giuridica corretta: Cloudflare agisce come TITOLARE AUTONOMO DEL TRATTAMENTO '
        '(independent data controller) per i dati tecnici dei visitatori (indirizzo IP, header HTTP, dati di routing) '
        'che elabora per proprie finalità di sicurezza della rete.',
        size=10, color=GRAY, space_after=3)
    add_para(doc,
        'Fonte: Privacy Policy di Cloudflare, Sez. 6: "Cloudflare is a data controller for the personal information '
        'collected from all categories of data subjects listed above" (End User = visitatori dei siti dei clienti Cloudflare).',
        italic=True, size=9, color=GRAY, space_after=4)

    cf_punti = [
        'L\'Associazione NON nomina Cloudflare come Responsabile del Trattamento ai sensi dell\'Art. 28 GDPR.',
        'Non è richiesta la stipula di un DPA con Cloudflare per i dati dei visitatori.',
        'Il trasferimento verso USA è lecito in base all\'EU-US Data Privacy Framework (Dec. adeguatezza CE 2023/1795 del 10/07/2023).',
        'Cloudflare va menzionata nell\'informativa privacy come terza parte titolare autonoma, con rinvio alla sua Privacy Policy.',
    ]
    for punto in cf_punti:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(punto)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    add_para(doc, '6b. Resend API — Invio email transazionali',
             bold=True, size=10, color=BLUE, space_before=8, space_after=2)
    add_para(doc,
        'Qualificazione: Responsabile del Trattamento ex Art. 28 GDPR. '
        'La funzione sendEmail() è presente nel codice ma NON è attualmente collegata ad alcun endpoint attivo. '
        'Quando questa funzionalità sarà attivata, occorrerà stipulare il DPA con Resend e aggiornare l\'informativa.',
        size=10, color=GRAY, space_after=6)

    add_para(doc, '6c. jsDelivr CDN (ProspectOne Sp. z o.o., Polonia — UE)',
             bold=True, size=10, color=BLUE, space_before=4, space_after=2)
    add_para(doc,
        'Qualificazione: titolare autonomo per la trasmissione tecnica dell\'IP al caricamento dei file CSS/JS '
        '(Tailwind, FontAwesome). Nessun cookie di profilazione. Trattamento interamente in UE.',
        size=10, color=GRAY, space_after=8)

    add_hline(doc)

    # ── 7. MISURE DI SICUREZZA ────────────────────────────────────────────────
    add_heading(doc, '7. MISURE DI SICUREZZA ADOTTATE')
    misure = [
        'Trasmissione HTTPS/TLS su tutta la rete Cloudflare.',
        'Hash IP crittografico: SHA-256 con salt fisso (renu-ip-salt-2026:), troncato a 16 caratteri hex, prefisso "sha2:". Operazione non reversibile.',
        'Consenso GDPR versione 2.0 registrato su entrambi i form, con versione testo e timestamp.',
        'Accesso admin protetto da token segreto cifrato su Cloudflare (mai nel codice sorgente).',
        'Audit log di tutte le operazioni amministrative con tracciabilità completa.',
        'Diritto all\'oblio automatizzato tramite pseudonimizzazione multi-tabella.',
        'Infrastruttura certificata ISO 27001 (Cloudflare Pages).',
    ]
    for m in misure:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(m)
        r.font.size = Pt(10); r.font.color.rgb = GRAY

    add_hline(doc)

    # ── 8. RICHIESTE AL DPO ───────────────────────────────────────────────────
    add_heading(doc, '8. RICHIESTE ALL\'AVV. CONTI — DPO')
    add_para(doc, 'Si chiede all\'Avv. Conti di:', size=10, color=GRAY, space_after=4)
    richieste = [
        'Validare la presente ricognizione e segnalare eventuali lacune o inesattezze.',
        'Verificare e controfirmare la bozza di contratto DPA Art. 28 GDPR allegata, da sottoscrivere con il Responsabile Tecnico.',
        'Confermare l\'impostazione giuridica adottata per Cloudflare (titolare autonomo, non responsabile nominabile ex Art. 28 GDPR).',
        'Indicare se la funzione sendEmail() tramite Resend, prima di essere attivata, richiede ulteriori adempimenti (DPA Resend, aggiornamento informativa, registro trattamenti).',
        'Valutare se il trattamento dei dati sanitari dei minori (storie, diagnosi, varianti genetiche) richiede una DPIA (Valutazione d\'Impatto ex Art. 35 GDPR).',
    ]
    for i, r_text in enumerate(richieste, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r_n = p.add_run(f'{i}.  ')
        r_n.bold = True; r_n.font.size = Pt(10); r_n.font.color.rgb = BLUE
        r_b = p.add_run(r_text)
        r_b.font.size = Pt(10); r_b.font.color.rgb = GRAY

    doc.add_paragraph()
    add_hline(doc)

    # Chiusura
    add_para(doc, 'Restiamo a disposizione per qualsiasi chiarimento o integrazione.', size=10, color=GRAY, space_before=6, space_after=2)
    add_para(doc, 'Cordiali saluti,', size=10, color=GRAY, space_after=10)

    add_field(doc, '', 'Stefania Rocca')
    add_field(doc, '', 'Presidente — Sindrome ReNU Italia APS')
    add_field(doc, '', 'presidenza@sindromerenu.it')
    doc.add_paragraph()
    add_para(doc, 'Data: _____________________        Firma: _____________________',
             size=10, color=GRAY, space_after=8)

    add_hline(doc)
    add_para(doc, 'Allegati: Contratto DPA Art. 28 GDPR — Responsabile Tecnico del Sito',
             italic=True, size=9, color=GRAY, space_after=4)

    path = '/home/user/webapp-renu/Lettera_Avv_Conti_DPO_GDPR.docx'
    doc.save(path)
    print(f'✅ Salvato: {path}')
    return path


# ── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    p1 = crea_dpa()
    p2 = crea_lettera()
    print(f'\n📄 Documenti generati:\n  {p1}\n  {p2}')
